import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import json

st.set_page_config(page_title="YouTube to Quiz Architect", layout="wide")
st.title("YouTube to Quiz Architect 🛠️")

# --- API KEY HANDLING ---
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    api_key = st.sidebar.text_input("Enter Gemini API Key:", type="password")

if api_key:
    genai.configure(api_key=api_key)
else:
    st.info("👉 Please ensure your Gemini API Key is set to activate the app.")
    st.stop()

# --- STEP 1: VIDEO URL INPUT ---
video_url = st.text_input("1. Paste YouTube URL here:", placeholder="https://www.youtube.com/watch?v=...")

if video_url:
    st.markdown("### Next Step: Get the Transcript")
    gemini_instruction = (
        f"Extract the complete caption/transcript text of this video: {video_url}\n\n"
        f"Format Requirements:\n"
        f"1. Put a new line at the end of each sentence.\n"
        f"2. Produce the entire final output inside a plain text code block so that it has a copy button."
    )
    st.code(gemini_instruction, language="text")

st.divider()

# --- STEP 3: TRANSCRIPT INPUT & PROCESSING ---
transcript_text = st.text_area("2. Paste your video transcript text below", height=300)

# --- DOCX COMPILER (The Pivot) ---
def generate_docx(url_str, questions_list):
    doc = Document()
    
    # Add URL
    doc.add_paragraph(url_str if url_str else "No URL Provided")
    doc.add_paragraph("") # Empty paragraph for spacing

    explanations = []
    
    for i, q in enumerate(questions_list):
        # Every .add_paragraph() creates a genuine hard return/paragraph break
        doc.add_paragraph(f"{i+1}) {q['question']}")
        
        for option in q['options']:
            doc.add_paragraph(option)
        
        doc.add_paragraph(f"ANSWER: {q['Answer']}")
        doc.add_paragraph(f"POINT: {q.get('points', 1)}")
        doc.add_paragraph("") # Space between questions
        
        explanations.append(q['explanation'])
        
    doc.add_paragraph("Question: Metadata")
    metadata_string = "**".join(explanations)
    doc.add_paragraph(metadata_string)
    
    # Save to a buffer so Streamlit can download it
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- AI GENERATION LOGIC ---
if st.button("Generate Questions from Transcript", type="primary"):
    if not transcript_text:
        st.warning("Please paste a transcript block.")
    else:
        with st.spinner("AI is building your questions..."):
            try:
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = f"""
                Analyze the content and generate 15 MCQs.
                Return ONLY valid JSON:
                {{
                  "questions": [
                    {{
                      "question": "...",
                      "options": ["A...", "B...", "C...", "D..."],
                      "Answer": "A",
                      "explanation": "...",
                      "points": 1
                    }}
                  ]
                }}
                TEXT: {transcript_text}
                """
                response = model.generate_content(prompt, generation_config={"response_mime_type": "application/json"})
                data = json.loads(response.text)
                st.session_state['quiz_data'] = data['questions']
                st.session_state['saved_url'] = video_url
            except Exception as e:
                st.error(f"Error: {e}")

# --- EDITABLE UI ---
if 'quiz_data' in st.session_state:
    final_compiled_questions = []
    for i, q in enumerate(st.session_state['quiz_data']):
        with st.expander(f"Edit Question {i+1}", expanded=True):
            edit_q = st.text_input("Question", value=q['question'], key=f"q_{i}")
            edit_opts = [st.text_input(f"Opt {j}", value=q['options'][j], key=f"o_{i}_{j}") for j in range(4)]
            edit_a = st.text_input("Answer", value=q['Answer'], key=f"a_{i}")
            edit_p = st.number_input("Point", value=int(q.get('points', 1)), key=f"p_{i}")
            edit_e = st.text_area("Explanation", value=q['explanation'], key=f"e_{i}")
            
            if st.checkbox("Include?", value=True, key=f"k_{i}"):
                final_compiled_questions.append({
                    "question": edit_q, "options": edit_opts, 
                    "Answer": edit_a, "points": edit_p, "explanation": edit_e
                })

    if final_compiled_questions:
        docx_data = generate_docx(st.session_state['saved_url'], final_compiled_questions)
        st.download_button(
            label="💾 Download Word Doc (.docx) for MS Forms",
            data=docx_data,
            file_name="quiz_for_forms.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
